from docx import Document
from docx.shared import Inches


def create_form_template():
    """Create a new Word template with placeholders for ID card information."""

    # Create a new document
    doc = Document()

    # Add title
    title = doc.add_heading("Input Form", 0)
    title.alignment = 1  # Center alignment

    # Add some space
    doc.add_paragraph()

    # Add form fields
    doc.add_paragraph(f"Name: {{{{NAME}}}}")
    doc.add_paragraph()

    doc.add_paragraph(f"Date of Birth: {{{{DOB}}}}")
    doc.add_paragraph()

    doc.add_paragraph(f"ID: {{{{ID}}}}")
    doc.add_paragraph()

    doc.add_paragraph(f"Address: {{{{ADDRESS}}}}")
    doc.add_paragraph()

    # Add gender checkboxes
    gender_para = doc.add_paragraph("Sex: M ")
    gender_para.add_run("{{GENDER_M}}")
    gender_para.add_run("  F ")
    gender_para.add_run("{{GENDER_F}}")

    doc.add_paragraph()
    doc.add_paragraph()

    # Add AI prediction info (optional, can be removed)
    doc.add_paragraph("{{AI_PREDICTION}}")

    # Save the template
    doc.save("Input Form_Template.docx")
    print("✅ Created new template: Input Form_Template.docx")


if __name__ == "__main__":
    create_form_template()
