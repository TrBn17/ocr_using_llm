import json
import os
import re
import shutil
import sys

import google.generativeai as genai
from docx import Document
from PIL import Image

from infra.config import GEMINI_API_KEY


def sanitize_filename(name):
    """Remove invalid characters from a string to use as a filename."""
    return re.sub(r'[\\/*?:"<>|]', "", name)


def docx_replace(doc, replacements):
    """
    Find and replace text in a .docx file, including headers, footers, and body.
    `replacements` is a dictionary of {"placeholder": "new_text"}.
    """
    # Replace in body
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                inline = p.runs
                # Replace strings and retain formatting
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, value)
                        inline[i].text = text

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace(cell, replacements)  # Recursive call for cells


def extract_and_create_document(image_path):
    """
    Extract data from ID card and create a new document with extracted information.
    """
    # Configure Gemini API
    api_key = GEMINI_API_KEY

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-2.0-flash")

    # Check if image exists
    if not os.path.exists(image_path):
        print(f"❌ Image file not found: {image_path}")
        return False

    # Ensure results directory exists
    if not os.path.exists("results"):
        os.makedirs("results")
        print("📁 Created 'results' directory.")

    # Find the template file
    template_paths = [
        "form/Input Form_Template.docx",
        "../form/Input Form_Template.docx",
        "form\\Input Form_Template.docx",
        "../form\\Input Form_Template.docx",
        "form/Input Form_Tmp.docx",
        "../form/Input Form_Tmp.docx",
        "form\\Input Form_Tmp.docx",
        "../form\\Input Form_Tmp.docx",
    ]
    template_path = None
    for path in template_paths:
        if os.path.exists(path):
            template_path = path
            break

    if not template_path:
        print(
            f"❌ Template file not found. Looking for 'Input Form_Template.docx' or 'Input Form_Tmp.docx'"
        )
        return False

    print(f"🔍 Processing image: {image_path}")

    try:
        # Open image
        image = Image.open(image_path)

        # Create enhanced prompt
        prompt = """
        Phân tích kỹ hình ảnh thẻ căn cước này và trích xuất thông tin theo định dạng JSON.
        Hãy chắc chắn rằng các trường sau được điền đầy đủ và chính xác:
        {
            "name": "Họ tên đầy đủ từ thẻ",
            "date_of_birth": "Ngày sinh (định dạng DD/MM/YYYY)", 
            "id": "Số CMND/CCCD",
            "address": "Địa chỉ đầy đủ",
            "predicted_gender": "Dự đoán giới tính từ tên (M/F/Unknown)",
            "gender_confidence": "Độ tin cậy dự đoán (high/medium/low)"
        }
        
        Hướng dẫn dự đoán giới tính:
        - Phân tích tên để đoán giới tính.
        - Tên Lào: Tên nam thường có vẻ mạnh mẽ, tên nữ thường nhẹ nhàng hơn.
        - Trả về "M" cho nam, "F" cho nữ, "Unknown" nếu không chắc.
        - Confidence: "high" nếu chắc chắn, "medium" nếu có thể, "low" nếu khó đoán.
        
        Nếu không thấy rõ trường nào, sử dụng giá trị "N/A".
        Chỉ trả về đối tượng JSON, không có bất kỳ văn bản giải thích hay định dạng markdown nào.
        """

        # Get response from Gemini
        response = model.generate_content([prompt, image])
        response_text = response.text.strip().replace("```json", "").replace("```", "")

        # Parse JSON
        try:
            data = json.loads(response_text)
        except json.JSONDecodeError:
            print(f"❌ Error parsing JSON from AI. Response was: {response_text}")
            return False

        print("✅ Extracted data:")
        for key, value in data.items():
            print(f"   {key}: {value}")

        # Handle gender determination automatically
        predicted_gender = data.get("predicted_gender", "Unknown").upper()
        gender_confidence = data.get("gender_confidence", "low").lower()

        print(
            f"\n🤖 AI predicts gender: {predicted_gender} (confidence: {gender_confidence})"
        )

        # Use AI prediction if available, otherwise set as UNKNOWN
        if predicted_gender in ["M", "F"]:
            gender = predicted_gender
            print(
                f"✅ Giới tính được xác định tự động: {'Nam' if gender == 'M' else 'Nữ'} ({gender})"
            )
        else:
            gender = "UNKNOWN"
            print("⚠️  Giới tính không thể xác định từ AI")

        # Define replacements for the template
        male_checkbox = "☐"
        female_checkbox = "☐"
        if gender == "M":
            male_checkbox = "☑"
        elif gender == "F":
            female_checkbox = "☑"

        # Create dictionary for replacements
        replacements = {
            "{{NAME}}": data.get("name", "N/A"),
            "{{DOB}}": data.get("date_of_birth", "N/A"),
            "{{ID}}": data.get("id", "N/A"),
            "{{ADDRESS}}": data.get("address", "N/A"),
            "{{GENDER_M}}": male_checkbox,
            "{{GENDER_F}}": female_checkbox,
            "{{AI_PREDICTION}}": f"AI Prediction: {predicted_gender} (confidence: {gender_confidence})",
        }

        # --- Create and Save Word Document ---

        # 1. Generate filename for the filled form
        output_filename_base = f"{data.get('id', 'NO_ID')}_{sanitize_filename(data.get('name', 'NO_NAME'))}"
        docx_filename = f"{output_filename_base}.docx"
        docx_path = os.path.join("results", docx_filename)

        # 2. Copy template to the new path
        shutil.copy2(template_path, docx_path)

        # 3. Open the new document and replace placeholders
        doc = Document(docx_path)
        docx_replace(doc, replacements)

        # 4. Save the final filled document
        doc.save(docx_path)

        print(f"📄 Form đã được điền và lưu: {docx_path}")
        print(f"🎉 Hoàn thành xử lý cho: {data.get('name', 'Unknown')}")

        return True

    except Exception as e:
        print(f"❌ An unexpected error occurred: {str(e)}")
        return False


def process_multiple_images():
    """Process multiple ID card images automatically."""
    print("🆔 Batch Processing Mode - Lao ID Card Extractor 🆔")
    print("=" * 60)

    # Search for images
    print("Searching for images in ./, ./data/, ../data/")
    image_dirs = [".", "data", "../data"]
    image_files = []

    for dir_path in image_dirs:
        if os.path.exists(dir_path):
            try:
                files = [
                    os.path.join(dir_path, f)
                    for f in os.listdir(dir_path)
                    if f.lower().endswith((".png", ".jpg", ".jpeg"))
                ]
                image_files.extend(files)
            except OSError as e:
                print(f"Could not read directory {dir_path}: {e}")

    if not image_files:
        print(
            "❌ No image files found. Please place images in the script's directory or a 'data' subfolder."
        )
        return

    print(f"\n📁 Found {len(image_files)} image(s):")
    for i, filename in enumerate(image_files, 1):
        print(f"  {i}. {os.path.basename(filename)} ({os.path.dirname(filename)})")

    # Ask user for selection
    print("\n🔄 Processing options:")
    print("  1. Process all images")
    print("  2. Select specific images")

    try:
        mode = input("\n👉 Choose processing mode (1 or 2): ").strip()

        selected_images = []
        if mode == "1":
            selected_images = image_files
            print(f"✅ Selected all {len(image_files)} images for processing")
        elif mode == "2":
            print(f"\nSelect images to process (1-{len(image_files)}):")
            print(
                "Enter numbers separated by commas (e.g., 1,3,5) or ranges (e.g., 1-3,5):"
            )
            selection = input("👉 Your selection: ").strip()

            # Parse selection
            for part in selection.split(","):
                part = part.strip()
                if "-" in part:
                    # Handle range
                    try:
                        start, end = map(int, part.split("-"))
                        for i in range(start, end + 1):
                            if 1 <= i <= len(image_files):
                                selected_images.append(image_files[i - 1])
                    except ValueError:
                        print(f"❌ Invalid range: {part}")
                else:
                    # Handle single number
                    try:
                        i = int(part)
                        if 1 <= i <= len(image_files):
                            selected_images.append(image_files[i - 1])
                    except ValueError:
                        print(f"❌ Invalid number: {part}")

            selected_images = list(set(selected_images))  # Remove duplicates
            print(f"✅ Selected {len(selected_images)} images for processing")
        else:
            print("❌ Invalid selection. Please enter 1 or 2.")
            return

        if not selected_images:
            print("❌ No valid images selected.")
            return

        # Process selected images
        successful = 0
        failed = 0

        print(f"\n🚀 Starting batch processing of {len(selected_images)} images...")
        print("=" * 60)

        for i, image_path in enumerate(selected_images, 1):
            print(
                f"\n📸 Processing image {i}/{len(selected_images)}: {os.path.basename(image_path)}"
            )
            print("-" * 40)

            if extract_and_create_document(image_path):
                successful += 1
                print(f"✅ Successfully processed {os.path.basename(image_path)}")
            else:
                failed += 1
                print(f"❌ Failed to process {os.path.basename(image_path)}")

        # Summary
        print("\n" + "=" * 60)
        print("📊 BATCH PROCESSING SUMMARY")
        print("=" * 60)
        print(f"✅ Successfully processed: {successful} files")
        print(f"❌ Failed to process: {failed} files")
        print(f"📁 Output directory: results/")

    except (KeyboardInterrupt, EOFError):
        print("\n\n❌ Operation cancelled by user. Exiting.")
    except Exception as e:
        print(f"\nAn error occurred: {e}")


def main():
    """Main function to run the script."""
    print("🆔 Lao ID Card to Document Extractor 🆔")
    print("=" * 50)

    if len(sys.argv) > 1:
        image_path = sys.argv[1]
        extract_and_create_document(image_path)
        return

    # Interactive mode with batch processing
    process_multiple_images()


if __name__ == "__main__":
    main()
